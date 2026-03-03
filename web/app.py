"""
논문 번역 웹 애플리케이션 백엔드
FastAPI + Claude API 기반 영문 학술 논문 → 한국어 번역 PDF 생성
"""
import asyncio
import json
import os
import sys
import uuid
import subprocess
from pathlib import Path
from datetime import datetime
from threading import Lock

import re

import anthropic
import pdfplumber
from fastapi import FastAPI, Form, UploadFile, File, BackgroundTasks, HTTPException
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse, JSONResponse
import uvicorn

# ── 경로 설정 ────────────────────────────────────────────────────────────────
WEB_DIR     = Path(__file__).parent
BASE_DIR    = WEB_DIR.parent / "paper-translation"
UPLOAD_DIR  = BASE_DIR / "input"
OUTPUT_DIR  = BASE_DIR / "output"
SCRIPTS_DIR = BASE_DIR / "scripts"
JOBS_FILE   = WEB_DIR / "jobs.json"
STATIC_DIR  = WEB_DIR / "static"

UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

# ── 도메인 매핑 ──────────────────────────────────────────────────────────────
DOMAIN_LABELS = {
    "cs":        "컴퓨터과학/AI/ML",
    "physics":   "물리학",
    "chemistry": "화학",
    "medicine":  "의학",
    "biology":   "생물학",
    "economics": "경제학",
    "general":   "일반 학술",
}

DOC_TYPES = {
    "academic":  "학술논문",
    "insurance": "보험약관",
    "general":   "일반내용",
}

# ── 스레드 안전 jobs.json 접근 ────────────────────────────────────────────────
_lock = Lock()

def load_jobs() -> dict:
    with _lock:
        if JOBS_FILE.exists():
            try:
                return json.loads(JOBS_FILE.read_text(encoding="utf-8"))
            except Exception:
                pass
    return {"jobs": []}

def save_jobs(data: dict):
    with _lock:
        JOBS_FILE.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")

def get_job(job_id: str) -> dict | None:
    for job in load_jobs()["jobs"]:
        if job["id"] == job_id:
            return job
    return None

def update_job(job_id: str, updates: dict):
    with _lock:
        if JOBS_FILE.exists():
            data = json.loads(JOBS_FILE.read_text(encoding="utf-8"))
        else:
            data = {"jobs": []}
        for job in data["jobs"]:
            if job["id"] == job_id:
                job.update(updates)
                break
        JOBS_FILE.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")

def detect_domain(filename: str) -> str:
    name = filename.lower()
    for d in DOMAIN_LABELS:
        if d == "general":
            continue
        if name.startswith(d + "_") or f"_{d}_" in name:
            return d
    return "general"


def detect_language(pages_text: list[str]) -> str:
    """PDF에서 추출한 텍스트로 언어를 감지한다. 'ja' 또는 'en' 반환."""
    # 앞쪽 5페이지 샘플로 판단
    sample = " ".join(pages_text[:5])
    if not sample.strip():
        return "en"

    ja_count = sum(
        1 for ch in sample
        if '\u3040' <= ch <= '\u309f'   # 히라가나
        or '\u30a0' <= ch <= '\u30ff'   # 카타카나
        or '\u4e00' <= ch <= '\u9fff'   # CJK 한자 (공통)
        or '\u3400' <= ch <= '\u4dbf'   # CJK 확장 A
    )
    total = sum(1 for ch in sample if not ch.isspace())
    if total == 0:
        return "en"

    ratio = ja_count / total
    return "ja" if ratio > 0.05 else "en"


LANG_LABELS = {"en": "영어", "ja": "일본어"}

# ── 취소 플래그 (job_id → bool) ──────────────────────────────────────────────
_cancel_flags: dict[str, bool] = {}


class TranslationCancelledError(Exception):
    pass


def _check_cancel(job_id: str):
    """취소 요청이 들어왔으면 예외 발생."""
    if _cancel_flags.get(job_id):
        raise TranslationCancelledError("사용자가 번역을 취소했습니다.")


# ── 번역 파이프라인 (동기) ───────────────────────────────────────────────────
def _get_translation_prompt(doc_type: str, domain: str, lang: str) -> str:
    """문서 유형별·언어별 번역 시스템 프롬프트 반환."""
    lang_label = LANG_LABELS.get(lang, "영어")
    is_ja = (lang == "ja")

    if doc_type == "insurance":
        if is_ja:
            lang_rules = (
                "- 일본 법률·보험 전문 용어는 한국 표준 용어로 번역하고 첫 등장 시 괄호 안에 일본어 원문 병기하세요\n"
                "- 일본어 경어·문어 표현(ます体·である体)은 자연스러운 한국어 문어체로 번역하세요\n"
                "- 고유명사(인물명·기관명)는 원문(일본어) 유지 후 필요시 한국어 추가\n"
            )
        else:
            lang_rules = "- 영문 전문 용어는 첫 등장 시 괄호 안에 원문 병기하세요\n"
        return (
            f"당신은 {lang_label} 보험약관·금융·법률 문서를 한국어로 번역하는 전문가입니다.\n\n"
            "번역 규칙:\n"
            "- 보험 업계 표준 용어를 사용하세요 (보험계약자, 피보험자, 보험금, 보험료 등)\n"
            "- 조항 번호는 원본 형식 그대로 유지하세요 (第1条, 제1조, Article 1 등)\n"
            "- 금액·비율·날짜·기간 등 수치를 정확히 번역하세요\n"
            "- 법적 면책 문구의 원문 의미를 정확하게 전달하세요\n"
            + lang_rules +
            "- 마크다운 형식 유지: # 제목, ## 섹션, ### 부제목, **굵게**, *기울임*\n"
            "- 번역된 텍스트만 출력하고 추가 설명은 하지 마세요\n"
        )

    elif doc_type == "general":
        if is_ja:
            lang_rules = (
                "- 일본어 특유 표현·관용구는 한국어로 의미를 살려 번역하세요\n"
                "- 고유명사(인물명·지명·작품명)는 원문(일본어) 또는 통용 한국어 표기를 사용하세요\n"
                "- 일본어 경어 표현은 적절한 한국어 문체로 번역하세요\n"
            )
        else:
            lang_rules = "- 고유명사는 원문 또는 통용 한국어 표기를 사용하세요\n"
        return (
            f"당신은 {lang_label} 도서·문서를 한국어로 번역하는 전문가입니다.\n\n"
            "번역 규칙:\n"
            "- 자연스럽고 읽기 쉬운 한국어로 번역하세요\n"
            + lang_rules +
            "- 직역보다 의미 전달을 우선시하세요\n"
            "- 마크다운 형식 유지: # 제목, ## 섹션, ### 부제목, **굵게**, *기울임*\n"
            "- 번역된 텍스트만 출력하고 추가 설명은 하지 마세요\n"
        )

    else:  # academic (기본값)
        domain_label = DOMAIN_LABELS.get(domain, "일반 학술")
        if is_ja:
            lang_rules = (
                "- 일본어 학술 문체(です・ます体, である体)를 자연스러운 한국어 학술 문체로 번역하세요\n"
                "- 일본어 고유 표현·관용구는 한국어로 의미를 살려 번역하세요\n"
                "- 저자명·기관명은 원문(일본어/영어) 유지 후 필요시 괄호 안에 한국어 추가\n"
            )
        else:
            lang_rules = "- 저자명, 기관명, 모델명은 원문 유지 (필요시 괄호 안에 한국어 추가)\n"
        return (
            f"당신은 {domain_label} 분야 {lang_label} 학술 논문을 한국어로 번역하는 전문가입니다.\n\n"
            "번역 규칙:\n"
            "- 학술 전문 용어를 정확하게 번역하세요\n"
            "- LaTeX 수식($...$, $$...$$)은 원문 그대로 유지하세요\n"
            + lang_rules +
            "- 마크다운 형식 유지: # 제목, ## 섹션, ### 부제목, **굵게**, *기울임*\n"
            "- 번역된 텍스트만 출력하고 추가 설명은 하지 마세요\n"
        )


def _get_summary_prompt(doc_type: str) -> str:
    """문서 유형별 요약 구조 프롬프트 반환."""
    if doc_type == "insurance":
        return (
            "요약문 구조 (마크다운 형식):\n"
            "# [보험 상품명 또는 약관 핵심을 한 줄로]\n\n"
            "## 주요 보장 내용\n"
            "## 보험금 지급 조건\n"
            "## 면책 사항\n"
            "## 납입 조건 및 보험 기간\n"
            "## 주요 특약\n"
        )
    elif doc_type == "general":
        return (
            "요약문 구조 (마크다운 형식):\n"
            "# [핵심 주제를 한 줄로]\n\n"
            "## 핵심 주제\n"
            "## 주요 내용\n"
            "## 핵심 포인트\n"
            "## 결론\n"
        )
    else:  # academic
        return (
            "요약문 구조 (마크다운 형식):\n"
            "# [논문의 핵심 주제를 한 줄로]\n\n"
            "## 연구 배경 및 문제 제기\n"
            "## 제안 방법 및 접근법\n"
            "## 주요 결과\n"
            "## 결론 및 기여\n"
        )


def _translate_chunk(client: anthropic.Anthropic, text: str, domain: str, lang: str, job_id: str = None, doc_type: str = "academic") -> str:
    system_instr = _get_translation_prompt(doc_type, domain, lang)
    prompt = (
        system_instr
        + "\n번역할 텍스트:\n"
        "---\n"
        f"{text}\n"
        "---\n\n"
        "한국어 번역:"
    )

    # 스트리밍으로 수신하면서 토큰마다 취소 플래그 확인 → 즉각 취소 가능
    parts: list[str] = []
    with client.messages.stream(
        model="claude-sonnet-4-6",
        max_tokens=4096,
        messages=[{"role": "user", "content": prompt}]
    ) as stream:
        for token in stream.text_stream:
            if job_id and _cancel_flags.get(job_id):
                raise TranslationCancelledError("사용자가 번역을 취소했습니다.")
            parts.append(token)
    return "".join(parts)


def _process_translation_sync(job_id: str):
    """백그라운드 번역 파이프라인 (동기)."""
    job = get_job(job_id)
    if not job:
        return

    try:
        input_path = Path(job["input_path"])
        paper_id   = job["paper_id"]
        output_dir = OUTPUT_DIR / paper_id
        output_dir.mkdir(parents=True, exist_ok=True)

        # API 키 확인
        if not os.environ.get("ANTHROPIC_API_KEY"):
            raise EnvironmentError(
                "ANTHROPIC_API_KEY 환경변수가 설정되지 않았습니다. "
                "서버 실행 전 set ANTHROPIC_API_KEY=sk-ant-... 으로 설정하세요."
            )

        # ── Step 1: PDF 텍스트 추출 ──────────────────────────────────────
        update_job(job_id, {"status": "processing", "current_step": "PDF 텍스트 추출 중", "progress": 5})

        pages_text = []
        with pdfplumber.open(str(input_path)) as pdf:
            total_pages = len(pdf.pages)
            update_job(job_id, {"page_count": total_pages})
            for i, page in enumerate(pdf.pages):
                _check_cancel(job_id)                          # ← 취소 체크
                text = page.extract_text() or ""
                pages_text.append(text)
                prog = 5 + int(10 * (i + 1) / total_pages)
                update_job(job_id, {"progress": prog, "current_step": f"PDF 파싱 ({i+1}/{total_pages} 페이지)"})

        # ── Step 1.5: 언어 감지 ──────────────────────────────────────────
        _check_cancel(job_id)
        lang = job.get("lang") or detect_language(pages_text)
        update_job(job_id, {"lang": lang, "current_step": f"언어 감지 완료: {LANG_LABELS.get(lang, lang)}"})

        # ── Step 2: 청크 분할 (~4000자 기준) ─────────────────────────────
        CHUNK_TARGET = 4000
        chunks = []
        current_chunk: list[str] = []
        current_len = 0

        for page_text in pages_text:
            if not page_text.strip():
                continue
            if current_len + len(page_text) > CHUNK_TARGET and current_chunk:
                chunks.append("\n\n".join(current_chunk))
                current_chunk = [page_text]
                current_len = len(page_text)
            else:
                current_chunk.append(page_text)
                current_len += len(page_text)

        if current_chunk:
            chunks.append("\n\n".join(current_chunk))

        if not chunks:
            raise ValueError(
                "PDF에서 텍스트를 추출할 수 없습니다.\n"
                "스캔 이미지 PDF이거나 텍스트 레이어가 없는 파일일 수 있습니다."
            )

        # ── Step 3: Claude API 번역 ───────────────────────────────────────
        client   = anthropic.Anthropic()
        domain   = detect_domain(job["filename"])
        doc_type = job.get("doc_type", "academic")
        lang_label = LANG_LABELS.get(lang, lang)
        total_chunks = len(chunks)
        translated_parts: list[str] = []

        for i, chunk in enumerate(chunks):
            _check_cancel(job_id)                              # ← 청크마다 취소 체크
            progress = 15 + int(70 * i / total_chunks)
            update_job(job_id, {
                "progress": progress,
                "current_step": f"{lang_label} 번역 중 ({i+1}/{total_chunks} 청크 / {total_pages} 페이지)"
            })
            translated = _translate_chunk(client, chunk, domain, lang, job_id, doc_type)
            translated_parts.append(translated)

        # ── Step 4: 마크다운 파일 생성 ───────────────────────────────────
        _check_cancel(job_id)
        update_job(job_id, {"current_step": "번역 결과 취합 중", "progress": 88})
        md_content = "\n\n---\n\n".join(translated_parts)
        md_path = output_dir / "translation_ko.md"
        md_path.write_text(md_content, encoding="utf-8")

        # ── Step 5: 마크다운 → PDF ────────────────────────────────────────
        update_job(job_id, {"current_step": "PDF 생성 중", "progress": 92})
        output_pdf = output_dir / f"{paper_id}_번역.pdf"
        result = subprocess.run(
            [
                sys.executable,
                str(SCRIPTS_DIR / "md_to_pdf.py"),
                "--input",  str(md_path),
                "--output", str(output_pdf),
                "--footer", job["filename"],
            ],
            capture_output=True,
            text=True,
            cwd=str(BASE_DIR),
        )

        if result.returncode != 0:
            stderr = result.stderr.strip() or "(stderr 없음)"
            raise RuntimeError(f"PDF 생성 실패:\n{stderr}")

        # ── Step 6: 마크다운 → DOCX ──────────────────────────────────────
        update_job(job_id, {"current_step": "Word 파일 생성 중", "progress": 96})
        docx_path = output_dir / f"{paper_id}_번역.docx"
        output_docx = None
        try:
            _generate_docx(md_path, docx_path)
            output_docx = str(docx_path)
        except Exception:
            pass  # DOCX 생성 실패해도 번역은 완료 처리

        update_job(job_id, {
            "status":       "completed",
            "current_step": "완료",
            "progress":     100,
            "completed_at": datetime.now().isoformat(),
            "output_pdf":   str(output_pdf),
            "output_docx":  output_docx,
        })

    except TranslationCancelledError:
        update_job(job_id, {
            "status":       "cancelled",
            "current_step": "취소됨",
            "completed_at": datetime.now().isoformat(),
        })

    except Exception as exc:
        update_job(job_id, {
            "status":       "failed",
            "current_step": "오류 발생",
            "error":        str(exc),
            "completed_at": datetime.now().isoformat(),
        })

    finally:
        _cancel_flags.pop(job_id, None)                        # ← 플래그 정리


async def process_translation(job_id: str):
    await asyncio.to_thread(_process_translation_sync, job_id)


# ── DOCX 생성 ────────────────────────────────────────────────────────────────
def _parse_table_row(line: str) -> list[str]:
    """파이프 테이블 행을 셀 목록으로 파싱."""
    return [c.strip() for c in line.strip().strip('|').split('|')]


def _add_inline_runs(paragraph, text: str):
    """마크다운 인라인 서식(**bold**, *italic*, `code`)을 Word run으로 변환."""
    from docx.shared import Pt
    pattern = re.compile(r'(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)')
    last = 0
    for m in pattern.finditer(text):
        if m.start() > last:
            paragraph.add_run(text[last:m.start()])
        full = m.group(0)
        if full.startswith('**'):
            run = paragraph.add_run(m.group(2))
            run.bold = True
        elif full.startswith('*'):
            run = paragraph.add_run(m.group(3))
            run.italic = True
        else:
            run = paragraph.add_run(m.group(4))
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        last = m.end()
    if last < len(text):
        paragraph.add_run(text[last:])


def _generate_docx(md_path: Path, docx_path: Path):
    """마크다운 파일을 Word(.docx) 파일로 변환."""
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    content = md_path.read_text(encoding="utf-8")
    lines = content.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 코드 블록
        if stripped.startswith("```"):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            p = doc.add_paragraph()
            run = p.add_run("\n".join(code_lines))
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            i += 1
            continue

        # 수평선 (청크 구분자)
        if re.match(r'^-{3,}$', stripped):
            doc.add_paragraph()
            i += 1
            continue

        # 헤더
        m = re.match(r'^(#{1,6})\s+(.*)', line)
        if m:
            level = min(len(m.group(1)), 3)
            doc.add_heading(m.group(2).strip(), level=level)
            i += 1
            continue

        # 테이블
        if stripped.startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            data_rows = [
                l for l in table_lines
                if not re.match(r'^\|[\s\-|:]+\|$', l.strip())
            ]
            if data_rows:
                parsed = [_parse_table_row(r) for r in data_rows]
                max_cols = max(len(r) for r in parsed)
                table = doc.add_table(rows=len(parsed), cols=max_cols)
                table.style = 'Table Grid'
                for r_idx, row_data in enumerate(parsed):
                    for c_idx in range(max_cols):
                        cell_text = row_data[c_idx] if c_idx < len(row_data) else ''
                        _add_inline_runs(
                            table.rows[r_idx].cells[c_idx].paragraphs[0], cell_text
                        )
            continue

        # 순서 없는 목록
        m = re.match(r'^\s*[-*+]\s+(.*)', line)
        if m:
            p = doc.add_paragraph(style='List Bullet')
            _add_inline_runs(p, m.group(1))
            i += 1
            continue

        # 순서 있는 목록
        m = re.match(r'^\s*\d+\.\s+(.*)', line)
        if m:
            p = doc.add_paragraph(style='List Number')
            _add_inline_runs(p, m.group(1))
            i += 1
            continue

        # 빈 줄
        if not stripped:
            i += 1
            continue

        # 일반 단락
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _add_inline_runs(p, line)
        i += 1

    doc.save(str(docx_path))


# ── 요약 파이프라인 ───────────────────────────────────────────────────────────
def _summarize_section(client: anthropic.Anthropic, text: str, domain: str, doc_type: str = "academic") -> str:
    """단일 섹션(청크)의 핵심 내용을 추출한다."""
    if doc_type == "insurance":
        doc_label = "보험약관"
    elif doc_type == "general":
        doc_label = "일반 문서"
    else:
        doc_label = DOMAIN_LABELS.get(domain, "일반 학술") + " 분야 논문"
    prompt = (
        f"다음은 {doc_label}의 한국어 번역 일부입니다.\n"
        "이 부분의 핵심 내용만 간결하게 추출해주세요.\n"
        "전문 용어와 수치 결과는 그대로 유지하고, 단락 형식으로 출력하세요.\n"
        "추가 설명 없이 핵심 내용만 출력하세요.\n\n"
        f"---\n{text}\n---\n\n핵심 내용:"
    )
    response = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=800,
        messages=[{"role": "user", "content": prompt}]
    )
    return response.content[0].text.strip()


def _generate_final_summary(client: anthropic.Anthropic, combined_points: str, domain: str, doc_type: str = "academic") -> str:
    """섹션별 핵심 내용 조각들로 구조화된 최종 요약을 생성한다."""
    if doc_type == "insurance":
        doc_label  = "보험약관"
        style_rule = "- 법률·금융 용어는 정확하게 사용하세요\n"
    elif doc_type == "general":
        doc_label  = "일반 문서"
        style_rule = "- 읽기 쉬운 자연스러운 문체로 작성하세요\n"
    else:
        doc_label  = DOMAIN_LABELS.get(domain, "일반 학술") + " 분야 학술 논문"
        style_rule = "- 중요한 수치와 통계는 반드시 포함하세요\n"
    structure = _get_summary_prompt(doc_type)
    prompt = (
        f"당신은 {doc_label} 요약 전문가입니다.\n\n"
        "다음은 문서의 각 섹션에서 추출한 핵심 내용입니다. "
        "이를 바탕으로 구조화된 한국어 요약문을 작성하세요.\n\n"
        + structure + "\n"
        "작성 규칙:\n"
        "- 각 섹션은 3-5 문장으로 핵심만 서술하세요\n"
        + style_rule +
        "- 전문 용어는 그대로 사용하세요\n"
        "- 추가 설명 없이 요약 본문만 출력하세요\n\n"
        "문서 핵심 내용:\n"
        f"---\n{combined_points}\n---\n\n"
        "한국어 요약:"
    )
    response = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=2048,
        messages=[{"role": "user", "content": prompt}]
    )
    return response.content[0].text.strip()


def _process_summary_sync(job_id: str):
    """백그라운드 요약 파이프라인 (동기)."""
    job = get_job(job_id)
    if not job:
        return

    try:
        paper_id   = job["paper_id"]
        output_dir = OUTPUT_DIR / paper_id
        md_path    = output_dir / "translation_ko.md"

        if not md_path.exists():
            raise FileNotFoundError(
                "번역 마크다운 파일을 찾을 수 없습니다. 먼저 번역을 완료하세요."
            )

        if not os.environ.get("ANTHROPIC_API_KEY"):
            raise EnvironmentError(
                "ANTHROPIC_API_KEY 환경변수가 설정되지 않았습니다."
            )

        # ── Step 1: 번역본 읽기 ───────────────────────────────────────
        update_job(job_id, {
            "summary_status":   "processing",
            "summary_step":     "번역본 읽는 중",
            "summary_progress": 5,
        })
        content = md_path.read_text(encoding="utf-8")

        # ── Step 2: 단락 기준 청크 분할 (~4000자) ────────────────────
        CHUNK_TARGET = 4000
        paragraphs = content.split("\n\n")
        chunks: list[str] = []
        current_chunk: list[str] = []
        current_len = 0

        for para in paragraphs:
            if not para.strip():
                continue
            if current_len + len(para) > CHUNK_TARGET and current_chunk:
                chunks.append("\n\n".join(current_chunk))
                current_chunk = [para]
                current_len   = len(para)
            else:
                current_chunk.append(para)
                current_len += len(para)

        if current_chunk:
            chunks.append("\n\n".join(current_chunk))

        if not chunks:
            raise ValueError("번역본 내용이 비어 있습니다.")

        # ── Step 3: 각 청크 핵심 내용 추출 ───────────────────────────
        client       = anthropic.Anthropic()
        domain       = detect_domain(job["filename"])
        doc_type     = job.get("doc_type", "academic")
        total_chunks = len(chunks)
        section_summaries: list[str] = []

        for i, chunk in enumerate(chunks):
            progress = 10 + int(55 * i / total_chunks)
            update_job(job_id, {
                "summary_step":     f"핵심 내용 추출 중 ({i+1}/{total_chunks})",
                "summary_progress": progress,
            })
            section_summaries.append(_summarize_section(client, chunk, domain, doc_type))

        # ── Step 4: 구조화된 최종 요약 생성 ──────────────────────────
        update_job(job_id, {"summary_step": "요약 정리 중", "summary_progress": 72})
        combined      = "\n\n".join(section_summaries)
        final_summary = _generate_final_summary(client, combined, domain, doc_type)

        # ── Step 5: 마크다운 저장 ──────────────────────────────────────
        update_job(job_id, {"summary_step": "요약 저장 중", "summary_progress": 85})
        summary_md = output_dir / "summary_ko.md"
        summary_md.write_text(final_summary, encoding="utf-8")

        # ── Step 6: PDF 생성 (font-boost=2) ───────────────────────────
        update_job(job_id, {"summary_step": "요약 PDF 생성 중", "summary_progress": 90})
        summary_pdf = output_dir / f"{paper_id}_요약.pdf"
        result = subprocess.run(
            [
                sys.executable,
                str(SCRIPTS_DIR / "md_to_pdf.py"),
                "--input",      str(summary_md),
                "--output",     str(summary_pdf),
                "--footer",     job["filename"],
                "--font-boost", "2",
            ],
            capture_output=True,
            text=True,
            cwd=str(BASE_DIR),
        )

        if result.returncode != 0:
            stderr = result.stderr.strip() or "(stderr 없음)"
            raise RuntimeError(f"요약 PDF 생성 실패:\n{stderr}")

        update_job(job_id, {
            "summary_status":   "completed",
            "summary_step":     "완료",
            "summary_progress": 100,
            "summary_pdf":      str(summary_pdf),
        })

    except Exception as exc:
        update_job(job_id, {
            "summary_status": "failed",
            "summary_step":   "오류 발생",
            "summary_error":  str(exc),
        })


async def process_summary(job_id: str):
    await asyncio.to_thread(_process_summary_sync, job_id)


# ── FastAPI 앱 ────────────────────────────────────────────────────────────────
app = FastAPI(title="논문 번역 시스템")


@app.post("/api/upload")
async def upload_file(
    file: UploadFile = File(...),
    doc_type: str = Form("academic"),
):
    """PDF 업로드만 수행. 번역은 /api/jobs/{id}/start 호출 시 시작."""
    if not file.filename.lower().endswith(".pdf"):
        raise HTTPException(400, "PDF 파일만 업로드 가능합니다.")

    job_id         = str(uuid.uuid4())[:8]
    safe_name      = file.filename.replace(" ", "_")
    paper_id       = f"{job_id}_{Path(safe_name).stem}"
    input_filename = f"{paper_id}.pdf"

    # 원본 PDF 저장
    input_path = UPLOAD_DIR / input_filename
    input_path.write_bytes(await file.read())

    # 출력 디렉터리 생성
    (OUTPUT_DIR / paper_id).mkdir(parents=True, exist_ok=True)

    # 페이지 수 및 언어 미리 파악 (앞 5페이지 샘플)
    page_count  = 0
    sample_text = []
    try:
        with pdfplumber.open(str(input_path)) as pdf:
            page_count = len(pdf.pages)
            for page in pdf.pages[:5]:
                sample_text.append(page.extract_text() or "")
    except Exception:
        pass

    lang = detect_language(sample_text)

    # Job 레코드 생성 (status: "ready" — 번역 대기)
    job = {
        "id":           job_id,
        "filename":     file.filename,
        "paper_id":     paper_id,
        "doc_type":     doc_type if doc_type in DOC_TYPES else "academic",
        "status":       "ready",
        "created_at":   datetime.now().isoformat(),
        "completed_at": None,
        "error":        None,
        "input_path":   str(input_path),
        "output_pdf":   None,
        "output_docx":  None,
        "progress":     0,
        "current_step": "번역 대기 중",
        "page_count":   page_count,
        "lang":         lang,
    }

    data = load_jobs()
    data["jobs"].insert(0, job)
    save_jobs(data)

    return {
        "job_id":     job_id,
        "filename":   file.filename,
        "page_count": page_count,
        "lang":       lang,
        "doc_type":   doc_type,
    }


@app.post("/api/jobs/{job_id}/start")
async def start_translation(job_id: str, background_tasks: BackgroundTasks):
    """번역 실행 버튼 클릭 시 호출 — 백그라운드 번역 시작."""
    job = get_job(job_id)
    if not job:
        raise HTTPException(404, "Job을 찾을 수 없습니다.")
    if job["status"] not in ("ready", "failed"):
        raise HTTPException(400, f"현재 상태({job['status']})에서는 시작할 수 없습니다.")

    update_job(job_id, {"status": "pending", "current_step": "대기 중", "progress": 0, "error": None})
    background_tasks.add_task(process_translation, job_id)
    return {"job_id": job_id, "status": "pending"}


@app.post("/api/jobs/{job_id}/cancel")
async def cancel_job(job_id: str):
    """번역 취소 요청 — 다음 청크 시작 전에 중단."""
    job = get_job(job_id)
    if not job:
        raise HTTPException(404, "Job을 찾을 수 없습니다.")
    if job["status"] not in ("pending", "processing"):
        raise HTTPException(400, f"취소할 수 없는 상태입니다: {job['status']}")

    _cancel_flags[job_id] = True
    update_job(job_id, {"current_step": "취소 요청됨..."})
    return {"job_id": job_id, "status": "cancelling"}


@app.get("/api/jobs")
async def get_all_jobs():
    return load_jobs()["jobs"]


@app.get("/api/jobs/{job_id}")
async def get_job_status(job_id: str):
    job = get_job(job_id)
    if not job:
        raise HTTPException(404, "Job을 찾을 수 없습니다.")
    return job


@app.get("/api/jobs/{job_id}/view-translated")
async def view_translated(job_id: str):
    job = get_job(job_id)
    if not job or job["status"] != "completed":
        raise HTTPException(404, "번역된 PDF가 없습니다.")
    pdf_path = Path(job["output_pdf"])
    if not pdf_path.exists():
        raise HTTPException(404, "파일을 찾을 수 없습니다.")
    return FileResponse(str(pdf_path), media_type="application/pdf")


@app.get("/api/jobs/{job_id}/view-original")
async def view_original(job_id: str):
    job = get_job(job_id)
    if not job:
        raise HTTPException(404, "Job을 찾을 수 없습니다.")
    input_path = Path(job["input_path"])
    if not input_path.exists():
        raise HTTPException(404, "원본 파일을 찾을 수 없습니다.")
    return FileResponse(str(input_path), media_type="application/pdf")


@app.post("/api/jobs/{job_id}/summarize")
async def start_summary(job_id: str, background_tasks: BackgroundTasks):
    """요약 생성 시작 — 번역 완료 후 호출 가능."""
    job = get_job(job_id)
    if not job:
        raise HTTPException(404, "Job을 찾을 수 없습니다.")
    if job["status"] != "completed":
        raise HTTPException(400, "번역이 완료된 후 요약을 생성할 수 있습니다.")
    if job.get("summary_status") in ("pending", "processing"):
        raise HTTPException(400, "요약이 이미 진행 중입니다.")

    update_job(job_id, {
        "summary_status":   "pending",
        "summary_step":     "요약 대기 중",
        "summary_progress": 0,
        "summary_error":    None,
        "summary_pdf":      None,
    })
    background_tasks.add_task(process_summary, job_id)
    return {"job_id": job_id, "summary_status": "pending"}


@app.get("/api/jobs/{job_id}/view-summary")
async def view_summary(job_id: str):
    """요약 PDF 뷰어."""
    job = get_job(job_id)
    if not job or job.get("summary_status") != "completed":
        raise HTTPException(404, "요약 PDF가 없습니다.")
    pdf_path = Path(job["summary_pdf"])
    if not pdf_path.exists():
        raise HTTPException(404, "파일을 찾을 수 없습니다.")
    return FileResponse(str(pdf_path), media_type="application/pdf")


@app.get("/api/jobs/{job_id}/download-summary")
async def download_summary(job_id: str):
    """요약 PDF 다운로드."""
    job = get_job(job_id)
    if not job or job.get("summary_status") != "completed":
        raise HTTPException(404, "요약 PDF가 없습니다.")
    pdf_path = Path(job["summary_pdf"])
    if not pdf_path.exists():
        raise HTTPException(404, "파일을 찾을 수 없습니다.")
    download_name = f"{Path(job['filename']).stem}_요약.pdf"
    return FileResponse(
        str(pdf_path),
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{download_name}"'},
    )


@app.get("/api/jobs/{job_id}/download")
async def download_translated(job_id: str):
    job = get_job(job_id)
    if not job or job["status"] != "completed":
        raise HTTPException(404, "번역된 PDF가 없습니다.")
    pdf_path = Path(job["output_pdf"])
    if not pdf_path.exists():
        raise HTTPException(404, "파일을 찾을 수 없습니다.")

    download_name = f"{Path(job['filename']).stem}_번역.pdf"
    return FileResponse(
        str(pdf_path),
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{download_name}"'},
    )


@app.get("/api/jobs/{job_id}/download-word")
async def download_word(job_id: str):
    """번역본 Word(.docx) 다운로드."""
    job = get_job(job_id)
    if not job or job["status"] != "completed":
        raise HTTPException(404, "번역된 Word 파일이 없습니다.")
    docx_path_str = job.get("output_docx")
    if not docx_path_str:
        raise HTTPException(404, "Word 파일이 생성되지 않았습니다.")
    docx_path = Path(docx_path_str)
    if not docx_path.exists():
        raise HTTPException(404, "파일을 찾을 수 없습니다.")
    download_name = f"{Path(job['filename']).stem}_번역.docx"
    return FileResponse(
        str(docx_path),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{download_name}"'},
    )


# ── 정적 파일 (마지막에 마운트) ───────────────────────────────────────────────
app.mount("/", StaticFiles(directory=str(STATIC_DIR), html=True), name="static")


if __name__ == "__main__":
    uvicorn.run("app:app", host="0.0.0.0", port=8000, reload=False)
